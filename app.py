from flask import Flask, render_template, request, redirect, send_file
import sqlite3
from docx import Document
import os

app = Flask(__name__)

DB = "database.db"

def init_db():
    conn = sqlite3.connect(DB)
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS pengajuan (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nama TEXT,
            alamat TEXT,
            prodi TEXT,
            keperluan TEXT
        )
    ''')
    conn.commit()
    conn.close()

init_db()

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/submit', methods=['POST'])
def submit():
    nama = request.form['nama']
    alamat = request.form['alamat']
    prodi = request.form['prodi']
    keperluan = request.form['keperluan']

    conn = sqlite3.connect(DB)
    c = conn.cursor()
    c.execute("INSERT INTO pengajuan (nama, alamat, prodi, keperluan) VALUES (?, ?, ?, ?)",
              (nama, alamat, prodi, keperluan))
    conn.commit()
    conn.close()

    return "Pengajuan berhasil dikirim!"

@app.route('/admin')
def admin():
    conn = sqlite3.connect(DB)
    c = conn.cursor()
    c.execute("SELECT * FROM pengajuan")
    data = c.fetchall()
    conn.close()
    return render_template('admin.html', data=data)

@app.route('/generate/<int:id>')
def generate(id):
    conn = sqlite3.connect(DB)
    c = conn.cursor()
    c.execute("SELECT * FROM pengajuan WHERE id=?", (id,))
    row = c.fetchone()
    conn.close()

    doc = Document()
    doc.add_heading("SURAT KETERANGAN", 0)
    doc.add_paragraph(f"Nama: {row[1]}")
    doc.add_paragraph(f"Alamat: {row[2]}")
    doc.add_paragraph(f"Program Studi: {row[3]}")
    doc.add_paragraph(f"Keperluan: {row[4]}")

    filename = f"generated/surat_{id}.docx"
    doc.save(filename)

    return send_file(filename, as_attachment=True)


if __name__ == '__main__':
    port = int(os.environ.get("PORT", 5000))
    app.run(host='0.0.0.0', port=port)